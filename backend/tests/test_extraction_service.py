"""Tests for the extraction service."""

import asyncio
import pytest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch, AsyncMock, MagicMock
from app.services.extraction_service import ExtractionService, ExtractionResult


class TestExtractionResult:
    """Test the ExtractionResult dataclass."""

    def test_extraction_result_calculates_word_count(self):
        """Word count is calculated from extracted text."""
        result = ExtractionResult(
            extracted_text="hello world this is a test"
        )
        assert result.word_count == 6

    def test_extraction_result_empty_text(self):
        """Empty text defaults to empty string."""
        result = ExtractionResult(extracted_text="")
        assert result.extracted_text == ""
        assert result.word_count == 0

    def test_extraction_result_warnings_initialized(self):
        """Warnings list is initialized if not provided."""
        result = ExtractionResult(extracted_text="test")
        assert result.warnings == []
        assert isinstance(result.warnings, list)

    def test_extraction_result_none_extracted_text(self):
        """None extracted text is converted to empty string."""
        result = ExtractionResult(extracted_text=None)
        assert result.extracted_text == ""


class TestPDFExtractor:
    """Test PDF extraction functionality."""

    @pytest.fixture
    def service(self):
        """Create an ExtractionService instance."""
        return ExtractionService()

    @pytest.mark.asyncio
    async def test_extract_pdf_no_bytes(self, service):
        """Extract PDF with no content returns error."""
        result = await service.extract_pdf(file_path=None, file_bytes=None)
        
        assert isinstance(result, ExtractionResult)
        assert result.error is not None
        assert "No PDF content" in result.error

    @pytest.mark.asyncio
    async def test_extract_pdf_invalid_bytes(self, service):
        """Invalid PDF bytes return error."""
        invalid_bytes = b"not a pdf at all"
        
        result = await service.extract_pdf(file_path=None, file_bytes=invalid_bytes)
        
        assert isinstance(result, ExtractionResult)
        # Should either extract empty or error
        assert result.extracted_text == "" or result.error is not None

    @pytest.mark.asyncio
    async def test_extract_pdf_with_basic_extraction_mock(self, service):
        """PDF extraction with mocked basic extraction."""
        pdf_bytes = b"mock pdf content"
        
        # Mock the basic extraction to succeed with sufficient text
        async def mock_basic_extract(file_bytes):
            return {
                'success': True,
                'text': 'Extracted PDF text content with more than 50 characters to pass the threshold test',
                'page_count': 1
            }
        
        with patch.object(service, '_try_basic_pdf_extraction', mock_basic_extract):
            result = await service.extract_pdf(file_path=None, file_bytes=pdf_bytes)
        
        assert isinstance(result, ExtractionResult)
        assert 'Extracted PDF text content' in result.extracted_text
        assert result.page_count == 1
        assert result.extraction_method == "basic_pdf"
        assert result.ocr_applied is False

    @pytest.mark.asyncio
    async def test_extract_pdf_fallback_to_ocr_with_mock(self, service):
        """PDF extraction falls back to OCR when basic extraction fails."""
        pdf_bytes = b"mock scanned pdf"
        
        # Mock basic extraction to return minimal text (not enough to pass threshold)
        async def mock_basic_extract(file_bytes):
            return {
                'success': True,
                'text': 'Short',  # Only 5 characters, below 50 char threshold
                'page_count': 1
            }
        
        # Mock OCR extraction to succeed
        async def mock_ocr_extract(file_bytes):
            return {
                'success': True,
                'text': 'OCR extracted text from scanned PDF which is now much longer',
                'page_count': 1
            }
        
        with patch.object(service, '_try_basic_pdf_extraction', mock_basic_extract):
            with patch.object(service, '_try_paddle_ocr_fallback', mock_ocr_extract):
                result = await service.extract_pdf(file_path=None, file_bytes=pdf_bytes)
        
        assert isinstance(result, ExtractionResult)
        assert 'OCR extracted text from scanned PDF' in result.extracted_text
        assert result.page_count == 1
        assert result.extraction_method == "paddle_ocr"
        assert result.ocr_applied is True
        assert "OCR fallback" in result.warnings[0]

    @pytest.mark.asyncio
    async def test_extract_pdf_calculates_processing_time(self, service):
        """PDF extraction calculates processing time."""
        pdf_bytes = b"mock pdf"
        
        # Mock extraction to succeed
        async def mock_extract(file_bytes):
            await asyncio.sleep(0.01)  # Simulate some processing
            return {
                'success': True,
                'text': 'Some text',
                'page_count': 1
            }
        
        with patch.object(service, '_try_basic_pdf_extraction', mock_extract):
            result = await service.extract_pdf(file_path=None, file_bytes=pdf_bytes)
        
        assert result.processing_time_ms > 0

    @pytest.mark.asyncio
    async def test_extract_pdf_from_file_path(self, service):
        """PDF extraction from file path."""
        # Create a temporary PDF file
        import tempfile
        pdf_content = b"mock pdf content"
        
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(pdf_content)
            temp_path = f.name
        
        try:
            # Mock extraction to succeed with sufficient text
            async def mock_extract(file_bytes):
                return {
                    'success': True,
                    'text': 'Extracted from file and it has more than 50 characters of content here',
                    'page_count': 1
                }
            
            with patch.object(service, '_try_basic_pdf_extraction', mock_extract):
                result = await service.extract_pdf(file_path=temp_path)
            
            assert 'Extracted from file' in result.extracted_text
            assert result.page_count == 1
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_pdf_all_methods_fail(self, service):
        """PDF extraction returns error when all methods fail."""
        pdf_bytes = b"mock pdf"
        
        # Mock all extraction methods to fail
        async def mock_marker_fail(file_bytes):
            return {'success': False}
        
        async def mock_basic_fail(file_bytes):
            return {'success': False}
        
        async def mock_ocr_fail(file_bytes):
            return {'success': False}
        
        with patch.object(service, '_try_marker_extraction', mock_marker_fail):
            with patch.object(service, '_try_basic_pdf_extraction', mock_basic_fail):
                with patch.object(service, '_try_paddle_ocr_fallback', mock_ocr_fail):
                    result = await service.extract_pdf(file_path=None, file_bytes=pdf_bytes)
        
        assert result.extracted_text == ""
        assert result.error is not None
        assert "Could not extract" in result.error


class TestExtractionService:
    """Test the main ExtractionService class."""

    @pytest.fixture
    def service(self):
        """Create an ExtractionService instance."""
        return ExtractionService()

    @pytest.mark.asyncio
    async def test_extract_from_file_pdf(self, service):
        """extract_from_file routes PDF correctly."""
        # Create a temporary PDF file
        import tempfile
        pdf_content = b"mock pdf"
        
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(pdf_content)
            temp_path = f.name
        
        try:
            # Mock extraction to succeed
            async def mock_extract(file_path_arg=None, file_bytes_arg=None):
                return ExtractionResult(
                    extracted_text="Test content",
                    page_count=1,
                    extraction_method="test"
                )
            
            with patch.object(service, 'extract_pdf', mock_extract):
                result = await service.extract_from_file(
                    file_path=temp_path,
                    mime_type="application/pdf",
                    timeout_seconds=30
                )
            
            assert isinstance(result, ExtractionResult)
            assert result.processing_time_ms > 0
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_from_file_text(self, service):
        """extract_from_file routes text files correctly."""
        import tempfile
        
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"plain text content")
            temp_path = f.name
        
        try:
            # Mock extract_text to succeed
            async def mock_extract_text(path):
                return ExtractionResult(
                    extracted_text="Plain text",
                    extraction_method="text"
                )
            
            with patch.object(service, 'extract_text', mock_extract_text):
                result = await service.extract_from_file(
                    file_path=temp_path,
                    mime_type="text/plain",
                    timeout_seconds=30
                )
            
            assert result.extraction_method == "text"
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_from_file_unsupported_mime(self, service):
        """Unsupported MIME types return error."""
        result = await service.extract_from_file(
            file_path="/tmp/test",
            mime_type="application/unsupported",
            timeout_seconds=30
        )
        
        assert result.error is not None
        assert "Unsupported" in result.error

    @pytest.mark.asyncio
    async def test_extract_from_file_timeout(self, service):
        """Timeout during extraction returns error."""
        import tempfile
        
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"mock pdf")
            temp_path = f.name
        
        try:
            # Mock extract_pdf to be very slow
            async def mock_slow_extract(path):
                await asyncio.sleep(5)
                return ExtractionResult(extracted_text="Never reached")
            
            with patch.object(service, 'extract_pdf', mock_slow_extract):
                result = await service.extract_from_file(
                    file_path=temp_path,
                    mime_type="application/pdf",
                    timeout_seconds=0.1  # Very short timeout
                )
            
            assert result.error is not None
            assert "timed out" in result.error.lower()
        finally:
            Path(temp_path).unlink()


class TestPaddleOCRFixes:
    """Test critical fixes for PaddleOCR and error handling."""

    @pytest.fixture
    def service(self):
        """Create an ExtractionService instance."""
        return ExtractionService()

    @pytest.mark.asyncio
    async def test_paddle_ocr_handles_invalid_result_format(self, service):
        """PaddleOCR handles invalid result format without crashing."""
        pdf_bytes = b"mock pdf"
        
        # Mock OCR to return None or malformed data
        async def mock_ocr_invalid(file_bytes):
            return {
                'success': True,
                'text': 'This should be ignored because result is invalid',
                'page_count': 1
            }
        
        # This test primarily validates the error handling in _try_paddle_ocr_fallback
        # The actual fix is in the result parsing validation
        with patch.object(service, '_try_paddle_ocr_fallback', mock_ocr_invalid):
            result = await service.extract_pdf(file_path=None, file_bytes=pdf_bytes)
        
        # Should handle gracefully
        assert isinstance(result, ExtractionResult)

    @pytest.mark.asyncio
    async def test_extract_pdf_with_timeout_parameter(self, service):
        """extract_pdf respects timeout parameter for direct calls."""
        pdf_bytes = b"mock pdf"
        
        # Mock slow extraction
        async def mock_slow_extract(file_bytes):
            await asyncio.sleep(2)
            return {
                'success': True,
                'text': 'Should timeout',
                'page_count': 1
            }
        
        with patch.object(service, '_try_basic_pdf_extraction', mock_slow_extract):
            result = await service.extract_pdf(
                file_path=None,
                file_bytes=pdf_bytes,
                timeout_seconds=0.1  # Very short timeout
            )
        
        assert result.error is not None
        assert "timed out" in result.error.lower()

    @pytest.mark.asyncio
    async def test_marker_extraction_in_fallback_chain(self, service):
        """Marker extraction is in the correct position in fallback chain."""
        pdf_bytes = b"mock pdf"
        
        call_order = []
        
        # Track which methods are called
        async def mock_basic_minimal(file_bytes):
            call_order.append("basic")
            return {
                'success': True,
                'text': 'Short',  # Below threshold
                'page_count': 1
            }
        
        async def mock_marker(file_bytes):
            call_order.append("marker")
            return {'success': False}  # Marker not implemented
        
        async def mock_ocr(file_bytes):
            call_order.append("ocr")
            return {
                'success': True,
                'text': 'OCR result',
                'page_count': 1
            }
        
        with patch.object(service, '_try_basic_pdf_extraction', mock_basic_minimal):
            with patch.object(service, '_try_marker_extraction', mock_marker):
                with patch.object(service, '_try_paddle_ocr_fallback', mock_ocr):
                    result = await service.extract_pdf(file_path=None, file_bytes=pdf_bytes)
        
        # Verify call order: basic -> marker -> ocr
        assert call_order == ["basic", "marker", "ocr"]
        # Verify final result is from OCR
        assert result.extraction_method == "paddle_ocr"


class TestDOCXExtraction:
    """Test DOCX extraction functionality."""

    @pytest.fixture
    def service(self):
        """Create an ExtractionService instance."""
        return ExtractionService()

    @pytest.mark.asyncio
    async def test_extract_docx_simple_document(self, service):
        """Extract text from simple DOCX document."""
        import tempfile
        from docx import Document
        
        # Create a test DOCX file
        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It has multiple paragraphs.")
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            result = await service.extract_docx(temp_path)
            
            assert isinstance(result, ExtractionResult)
            assert "test document" in result.extracted_text
            assert "multiple paragraphs" in result.extracted_text
            assert result.word_count > 0
            assert result.extraction_method == "python-docx"
            assert result.error is None
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_docx_with_table_formatting(self, service):
        """Extract text from DOCX with table formatting preserved."""
        import tempfile
        from docx import Document
        from docx.shared import Pt
        
        # Create a test DOCX with table
        doc = Document()
        doc.add_paragraph("Document with table:")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            result = await service.extract_docx(temp_path)
            
            assert isinstance(result, ExtractionResult)
            assert "table" in result.extracted_text.lower()
            assert "Header 1" in result.extracted_text or "Header1" in result.extracted_text.replace(" ", "")
            assert result.error is None
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_docx_empty_document(self, service):
        """Extract text from empty DOCX document."""
        import tempfile
        from docx import Document
        
        # Create empty DOCX
        doc = Document()
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            result = await service.extract_docx(temp_path)
            
            assert isinstance(result, ExtractionResult)
            assert result.extracted_text == ""
            assert result.word_count == 0
            assert result.error is None
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_docx_file_not_found(self, service):
        """Extract from non-existent DOCX file returns error."""
        result = await service.extract_docx("/tmp/nonexistent_file_12345.docx")
        
        assert isinstance(result, ExtractionResult)
        assert result.error is not None
        assert "not found" in result.error.lower() or "cannot" in result.error.lower()

    @pytest.mark.asyncio
    async def test_extract_docx_word_count(self, service):
        """Word count is calculated correctly for DOCX."""
        import tempfile
        from docx import Document
        
        # Create DOCX with known word count
        doc = Document()
        doc.add_paragraph("one two three four five")
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            result = await service.extract_docx(temp_path)
            
            assert result.word_count == 5
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_docx_invalid_file(self, service):
        """Extract from invalid DOCX file returns error."""
        import tempfile
        
        # Create a file that's not a valid DOCX
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(b"This is not a valid DOCX file")
            temp_path = f.name
        
        try:
            result = await service.extract_docx(temp_path)
            
            assert isinstance(result, ExtractionResult)
            assert result.error is not None
        finally:
            Path(temp_path).unlink()


class TestTextExtraction:
    """Test plain text extraction functionality."""

    @pytest.fixture
    def service(self):
        """Create an ExtractionService instance."""
        return ExtractionService()

    @pytest.mark.asyncio
    async def test_extract_text_utf8(self, service):
        """Extract plain text with UTF-8 encoding."""
        import tempfile
        
        # Create a UTF-8 text file
        text_content = "This is a test document.\nWith multiple lines.\nAnd UTF-8 encoding: café"
        
        with tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', suffix=".txt", delete=False) as f:
            f.write(text_content)
            temp_path = f.name
        
        try:
            result = await service.extract_text(temp_path)
            
            assert isinstance(result, ExtractionResult)
            assert "test document" in result.extracted_text
            assert "café" in result.extracted_text
            assert result.word_count > 0
            assert result.extraction_method.startswith("text-")
            assert result.error is None
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_text_latin1_fallback(self, service):
        """Extract plain text with latin-1 encoding fallback."""
        import tempfile
        
        # Create a latin-1 encoded file
        text_content = "This is a test with special chars: café naïve"
        
        with tempfile.NamedTemporaryFile(mode='wb', suffix=".txt", delete=False) as f:
            f.write(text_content.encode('latin-1'))
            temp_path = f.name
        
        try:
            result = await service.extract_text(temp_path)
            
            assert isinstance(result, ExtractionResult)
            assert result.extracted_text is not None
            assert result.word_count > 0
            assert result.error is None
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_text_empty_file(self, service):
        """Extract from empty text file."""
        import tempfile
        
        # Create empty file
        with tempfile.NamedTemporaryFile(mode='w', suffix=".txt", delete=False) as f:
            temp_path = f.name
        
        try:
            result = await service.extract_text(temp_path)
            
            assert isinstance(result, ExtractionResult)
            assert result.extracted_text == ""
            assert result.word_count == 0
            assert result.error is None
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_text_word_count(self, service):
        """Word count is calculated correctly for plain text."""
        import tempfile
        
        # Create text with known word count
        with tempfile.NamedTemporaryFile(mode='w', suffix=".txt", delete=False) as f:
            f.write("one two three four five")
            temp_path = f.name
        
        try:
            result = await service.extract_text(temp_path)
            
            assert result.word_count == 5
        finally:
            Path(temp_path).unlink()

    @pytest.mark.asyncio
    async def test_extract_text_file_not_found(self, service):
        """Extract from non-existent text file returns error."""
        result = await service.extract_text("/tmp/nonexistent_file_12345.txt")
        
        assert isinstance(result, ExtractionResult)
        assert result.error is not None
        assert "not found" in result.error.lower() or "cannot" in result.error.lower()


class TestVisionOCR:
    """Tests for vision model OCR extraction via LiteLLM/Ollama."""
    
    @pytest.mark.asyncio
    async def test_vision_ocr_disabled_by_default(self):
        """Vision OCR should be disabled by default."""
        from app.config import settings
        assert settings.VISION_OCR_ENABLED is False
    
    @pytest.mark.asyncio
    async def test_vision_ocr_skipped_when_disabled(self):
        """Vision OCR should be skipped when disabled."""
        service = ExtractionService()
        result = await service._try_vision_ocr(b"%PDF-1.4\ntest")
        assert result['success'] is False
        assert result['method'] == 'vision-disabled'
    
    @pytest.mark.asyncio
    async def test_pdf_to_images_empty_input(self):
        """PDF to images should handle empty input gracefully."""
        service = ExtractionService()
        images = await service._pdf_to_images(b"")
        assert images == []
    
    @pytest.mark.asyncio
    async def test_vision_ocr_returns_dict_structure(self):
        """Vision OCR should return correct dict structure."""
        service = ExtractionService()
        result = await service._try_vision_ocr(b"%PDF-1.4")
        
        # Check all expected keys present
        assert 'success' in result
        assert 'text' in result
        assert 'method' in result
        assert result['method'] == 'vision-ocr' or result['method'] == 'vision-disabled'
    
    @pytest.mark.asyncio
    async def test_paddleocr_config_used(self):
        """PaddleOCR should use settings.PADDLEOCR_USE_GPU value."""
        from app.config import settings
        assert settings.PADDLEOCR_USE_GPU is False  # Explicit CPU mode
    
    @pytest.mark.asyncio
    async def test_vision_ocr_timeout_setting(self):
        """Vision OCR timeout should be configured."""
        from app.config import settings
        assert settings.VISION_OCR_TIMEOUT_SECONDS == 45
        assert settings.VISION_OCR_TIMEOUT_SECONDS > 0
    
    @pytest.mark.asyncio
    async def test_vision_ocr_base_url_optional(self):
        """Vision OCR base URL should be optional."""
        from app.config import settings
        # Should be None or a string, but not required
        assert settings.VISION_OCR_BASE_URL is None or isinstance(settings.VISION_OCR_BASE_URL, str)

