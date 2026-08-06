"""Text extraction service for various file types."""

import logging
import asyncio
import time
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Callable
from pathlib import Path
from io import BytesIO

from app.config import settings

logger = logging.getLogger(__name__)

# Type for extraction function
ExtractionFunc = Callable[[bytes], Dict[str, any]]


@dataclass
class ExtractionResult:
    """Result of text extraction."""
    extracted_text: str
    page_count: Optional[int] = None
    word_count: int = 0
    extraction_method: str = "unknown"
    ocr_applied: bool = False
    processing_time_ms: float = 0.0
    warnings: List[str] = field(default_factory=list)
    error: Optional[str] = None
    
    def __post_init__(self):
        if not self.extracted_text:
            self.extracted_text = ""
        self.word_count = len(self.extracted_text.split()) if self.extracted_text else 0


class ExtractionService:
    """Service for extracting text from various file types."""
    
    async def extract_from_file(
        self,
        file_path: str,
        mime_type: str,
        timeout_seconds: int = 30
    ) -> ExtractionResult:
        """Main entry point - routes to appropriate extractor."""
        start_time = time.time()
        
        try:
            if mime_type == "application/pdf":
                result = await asyncio.wait_for(
                    self.extract_pdf(file_path),
                    timeout=timeout_seconds
                )
            elif mime_type.startswith("text/"):
                result = await self.extract_text(file_path)
            elif "word" in mime_type or "document" in mime_type:
                result = await self.extract_docx(file_path)
            else:
                return ExtractionResult(
                    extracted_text="",
                    error=f"Unsupported MIME type: {mime_type}"
                )
            
            result.processing_time_ms = (time.time() - start_time) * 1000
            return result
            
        except asyncio.TimeoutError:
            logger.warning(f"Extraction timed out for {file_path}")
            return ExtractionResult(
                extracted_text="",
                error=f"Extraction timed out after {timeout_seconds} seconds",
                processing_time_ms=(time.time() - start_time) * 1000
            )
        except Exception as e:
            logger.error(f"Extraction failed: {e}")
            return ExtractionResult(
                extracted_text="",
                error=str(e),
                processing_time_ms=(time.time() - start_time) * 1000
            )
    
    async def extract_pdf(self, file_path: str = None, file_bytes: bytes = None, timeout_seconds: int = 30) -> ExtractionResult:
        """Extract text from PDF using basic libraries, Marker, and PaddleOCR fallbacks.
        
        Args:
            file_path: Path to PDF file to read
            file_bytes: PDF content as bytes
            timeout_seconds: Timeout for extraction in seconds (default 30)
        
        Returns:
            ExtractionResult with extracted text and metadata
        """
        start_time = time.time()
        
        async def _do_extract():
            try:
                # Read file if path provided
                if file_path and not file_bytes:
                    with open(file_path, 'rb') as f:
                        file_bytes_local = f.read()
                else:
                    file_bytes_local = file_bytes
                
                if not file_bytes_local:
                    return ExtractionResult(
                        extracted_text="",
                        error="No PDF content provided"
                    )
                
                # Try basic PDF text extraction first (most reliable for modern PDFs)
                basic_result = await self._try_basic_pdf_extraction(file_bytes_local)
                if basic_result and basic_result.get('success'):
                    text = basic_result.get('text', '')
                    if len(text.strip()) > 50:  # If we got meaningful text
                        return ExtractionResult(
                            extracted_text=text,
                            page_count=basic_result.get('page_count'),
                            extraction_method="basic_pdf",
                            ocr_applied=False,
                            processing_time_ms=(time.time() - start_time) * 1000
                        )
                
                # Fall back to Marker extraction if configured (high-quality extraction, requires GPU)
                # This is intentionally last before OCR as it requires more complex setup
                logger.info("Basic extraction got minimal text, attempting Marker extraction")
                marker_result = await self._try_marker_extraction(file_bytes_local)
                
                if marker_result and marker_result.get('success'):
                    return ExtractionResult(
                        extracted_text=marker_result.get('text', ''),
                        page_count=marker_result.get('page_count'),
                        extraction_method="marker",
                        ocr_applied=False,
                        warnings=["Used Marker extraction for complex PDF"],
                        processing_time_ms=(time.time() - start_time) * 1000
                    )
                
                # Fall back to PaddleOCR for scanned PDFs
                logger.info("Marker extraction unavailable, attempting PaddleOCR fallback")
                ocr_result = await self._try_paddle_ocr_fallback(file_bytes_local)
                
                if ocr_result and ocr_result.get('success'):
                    return ExtractionResult(
                        extracted_text=ocr_result.get('text', ''),
                        page_count=ocr_result.get('page_count'),
                        extraction_method="paddle_ocr",
                        ocr_applied=True,
                        warnings=["Used OCR fallback for scanned PDF"],
                        processing_time_ms=(time.time() - start_time) * 1000
                    )
                
                return ExtractionResult(
                    extracted_text="",
                    error="Could not extract text from PDF with any available method",
                    processing_time_ms=(time.time() - start_time) * 1000
                )
                
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}", exc_info=True)
                return ExtractionResult(
                    extracted_text="",
                    error=str(e),
                    processing_time_ms=(time.time() - start_time) * 1000
                )
        
        try:
            # Apply timeout to the extraction
            result = await asyncio.wait_for(_do_extract(), timeout=timeout_seconds)
            return result
            
        except asyncio.TimeoutError:
            logger.warning(f"PDF extraction timed out after {timeout_seconds} seconds")
            return ExtractionResult(
                extracted_text="",
                error=f"PDF extraction timed out after {timeout_seconds} seconds",
                processing_time_ms=(time.time() - start_time) * 1000
            )
    
    async def _try_marker_extraction(self, file_bytes: bytes) -> Optional[Dict]:
        """Try to extract text from PDF using Marker library.
        
        Note: Marker extraction is a future enhancement. The marker-pdf library
        is complex and requires GPU/ML setup with significant dependencies.
        This method is intentionally placed last in the fallback chain before OCR,
        allowing for implementation when Marker is properly installed and configured.
        
        Returns:
            Dict with 'success' (bool) key, or None on error
        """
        try:
            # Marker implementation can be added here when dependencies are available
            # For now, signal that it's not yet implemented
            logger.debug("Marker extraction not yet implemented - skipping to OCR fallback")
            return {'success': False}
        except Exception as e:
            logger.debug(f"Marker extraction error: {e}")
            return {'success': False}
    
    async def _try_basic_pdf_extraction(self, file_bytes: bytes) -> Optional[Dict]:
        """Try basic PDF text extraction using available libraries."""
        try:
            loop = asyncio.get_event_loop()
            
            def _extract():
                try:
                    # Try PyMuPDF (fitz) first
                    try:
                        import fitz
                        doc = fitz.open(stream=file_bytes, filetype="pdf")
                        text = ""
                        for page in doc:
                            text += page.get_text()
                        return {
                            'success': True,
                            'text': text,
                            'page_count': len(doc)
                        }
                    except ImportError:
                        pass
                    
                    # Try pypdf
                    try:
                        from pypdf import PdfReader
                        reader = PdfReader(BytesIO(file_bytes))
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text()
                        return {
                            'success': True,
                            'text': text,
                            'page_count': len(reader.pages)
                        }
                    except ImportError:
                        pass
                    
                    # Try pdfplumber
                    try:
                        import pdfplumber
                        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                            text = ""
                            for page in pdf.pages:
                                text += page.extract_text() or ""
                            return {
                                'success': True,
                                'text': text,
                                'page_count': len(pdf.pages)
                            }
                    except ImportError:
                        pass
                    
                    return {'success': False}
                    
                except Exception as e:
                    logger.debug(f"Basic PDF extraction error: {e}")
                    return {'success': False}
            
            result = await loop.run_in_executor(None, _extract)
            return result
            
        except Exception as e:
            logger.debug(f"Basic extraction failed: {e}")
            return {'success': False}
    
    async def _pdf_to_images(self, file_bytes: bytes, dpi: int = 150) -> List[bytes]:
        """Convert PDF pages to PNG images for vision model processing.
        
        Args:
            file_bytes: PDF file as bytes
            dpi: Resolution for rendering (150 = good quality/size balance)
            
        Returns:
            List of PNG image bytes (one per page)
        """
        try:
            import fitz  # PyMuPDF (already in requirements)
            
            loop = asyncio.get_event_loop()
            
            def _convert():
                images = []
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    # Render page to pixmap at specified DPI
                    mat = fitz.Matrix(dpi / 72, dpi / 72)  # 72 DPI = default
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PNG bytes
                    img_bytes = pix.tobytes("png")
                    images.append(img_bytes)
                    
                doc.close()
                logger.info(f"Converted {len(images)} pages to images at {dpi} DPI")
                return images
            
            # Run in thread pool to avoid blocking
            images = await loop.run_in_executor(None, _convert)
            return images
            
        except Exception as e:
            logger.error(f"Failed to convert PDF to images: {e}", exc_info=True)
            return []
    
    async def _try_vision_ocr(self, file_bytes: bytes) -> Dict:
        """Extract text using LiteLLM vision model via Ollama.
        
        Args:
            file_bytes: PDF file as bytes
            
        Returns:
            Dict with 'success', 'text', 'method', 'page_count', 'warnings'
        """
        from app.config import settings
        
        if not settings.VISION_OCR_ENABLED:
            return {'success': False, 'text': '', 'method': 'vision-disabled'}
        
        try:
            from litellm import completion
            import base64
            
            # Convert PDF to images
            images = await self._pdf_to_images(file_bytes)
            if not images:
                return {
                    'success': False, 
                    'text': '', 
                    'method': 'vision-ocr',
                    'warnings': ['Failed to convert PDF to images']
                }
            
            # Process each page with vision model
            extracted_pages = []
            warnings = []
            
            for idx, img_bytes in enumerate(images):
                try:
                    # Encode image as base64
                    img_b64 = base64.b64encode(img_bytes).decode('utf-8')
                    
                    # Prepare kwargs for LiteLLM
                    completion_kwargs = {
                        'model': settings.VISION_OCR_MODEL,
                        'messages': [{
                            "role": "user",
                            "content": [
                                {"type": "text", "text": settings.VISION_OCR_PROMPT},
                                {"type": "image_url", "image_url": {
                                    "url": f"data:image/png;base64,{img_b64}"
                                }}
                            ]
                        }],
                        'timeout': settings.VISION_OCR_TIMEOUT_SECONDS
                    }
                    
                    # Add base_url if configured (for Ollama service)
                    if settings.VISION_OCR_BASE_URL:
                        completion_kwargs['base_url'] = settings.VISION_OCR_BASE_URL
                    
                    # Call vision model via LiteLLM
                    loop = asyncio.get_event_loop()
                    response = await asyncio.wait_for(
                        loop.run_in_executor(None, lambda: completion(**completion_kwargs)),
                        timeout=settings.VISION_OCR_TIMEOUT_SECONDS + 5
                    )
                    
                    # Extract text from response
                    page_text = response.choices[0].message.content
                    extracted_pages.append(f"--- Page {idx + 1} ---\n{page_text}")
                    logger.info(f"Vision OCR extracted {len(page_text)} chars from page {idx + 1}")
                    
                except asyncio.TimeoutError:
                    warnings.append(f"Page {idx + 1} timed out after {settings.VISION_OCR_TIMEOUT_SECONDS}s")
                    logger.warning(f"Vision OCR timeout on page {idx + 1}")
                except Exception as e:
                    warnings.append(f"Page {idx + 1} failed: {str(e)}")
                    logger.error(f"Vision OCR error on page {idx + 1}: {e}", exc_info=True)
            
            # Combine all pages
            full_text = "\n\n".join(extracted_pages)
            word_count = len(full_text.split())
            
            # Consider failure if very little text and many errors
            if word_count < 10 and len(warnings) >= len(images) * 0.5:
                return {
                    'success': False,
                    'text': full_text,
                    'method': 'vision-ocr',
                    'page_count': len(images),
                    'warnings': warnings
                }
            
            logger.info(f"Vision OCR extracted {word_count} words from {len(images)} pages")
            return {
                'success': True,
                'text': full_text,
                'method': 'vision-ocr',
                'page_count': len(images),
                'warnings': warnings if warnings else None
            }
            
        except Exception as e:
            logger.error(f"Vision OCR failed: {e}", exc_info=True)
            return {
                'success': False,
                'text': '',
                'method': 'vision-ocr',
                'warnings': [f"Vision OCR error: {str(e)}"]
            }
    
    async def _try_paddle_ocr_fallback(self, file_bytes: bytes) -> Dict:
        """Extract text from PDF using PaddleOCR as fallback for scanned PDFs."""
        try:
            loop = asyncio.get_event_loop()
            
            def _paddle_extract():
                try:
                    from paddleocr import PaddleOCR
                    import fitz
                    import cv2
                    import numpy as np
                    
                    # Initialize PaddleOCR
                    ocr = PaddleOCR(
                        use_angle_cls=True, 
                        lang='en',
                        use_gpu=settings.PADDLEOCR_USE_GPU,  # Explicit CPU-only
                        show_log=False  # Reduce noise in logs
                    )
                    
                    # Convert PDF to images
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    
                    all_text = ""
                    for page_num, page in enumerate(doc):
                        # Render page to image
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        # Convert PNG bytes to numpy array (PaddleOCR expects numpy array, not raw bytes)
                        png_bytes = pix.tobytes('png')
                        img_array = cv2.imdecode(
                            np.frombuffer(png_bytes, np.uint8), 
                            cv2.IMREAD_COLOR
                        )
                        
                        # Run OCR with numpy array
                        results = ocr.ocr(img_array, cls=True)
                        
                        # Extract text from results with proper error handling
                        if results and isinstance(results, list):
                            for line in results:
                                if line and isinstance(line, list):
                                    for word_info in line:
                                        # Validate word_info structure before accessing nested indices
                                        if (isinstance(word_info, (list, tuple)) and 
                                            len(word_info) > 1 and 
                                            isinstance(word_info[1], (list, tuple)) and 
                                            len(word_info[1]) > 0):
                                            all_text += str(word_info[1][0]) + " "
                                    all_text += "\n"
                    
                    return {
                        'success': True,
                        'text': all_text,
                        'page_count': len(doc)
                    }
                except ImportError as e:
                    logger.debug(f"PaddleOCR or dependencies not available: {e}")
                    return {'success': False}
                except Exception as e:
                    logger.error(f"PaddleOCR extraction failed: {e}", exc_info=True)
                    return {'success': False}
            
            result = await loop.run_in_executor(None, _paddle_extract)
            return result
            
        except Exception as e:
            logger.error(f"PaddleOCR fallback error: {e}")
            return {'success': False}
    
    async def extract_docx(self, file_path: str) -> ExtractionResult:
        """Extract text from DOCX using python-docx.
        
        Args:
            file_path: Path to DOCX file to read
        
        Returns:
            ExtractionResult with extracted text and metadata
        """
        start_time = time.time()
        
        try:
            from docx import Document
            
            # Open and read DOCX
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            
            # Extract text from tables (preserve table formatting)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        row_cells.append(cell.text)
                    text_parts.append(" | ".join(row_cells))
            
            extracted_text = "\n".join(text_parts)
            word_count = len(extracted_text.split()) if extracted_text else 0
            
            logger.info(f"Extracted {word_count} words from DOCX {file_path}")
            
            return ExtractionResult(
                extracted_text=extracted_text,
                word_count=word_count,
                extraction_method="python-docx",
                processing_time_ms=(time.time() - start_time) * 1000
            )
            
        except FileNotFoundError:
            logger.error(f"DOCX file not found: {file_path}")
            return ExtractionResult(
                extracted_text="",
                error=f"DOCX file not found: {file_path}",
                processing_time_ms=(time.time() - start_time) * 1000
            )
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}", exc_info=True)
            return ExtractionResult(
                extracted_text="",
                error=f"Could not extract text from DOCX: {str(e)}",
                processing_time_ms=(time.time() - start_time) * 1000
            )
    
    async def extract_text(self, file_path: str) -> ExtractionResult:
        """Extract plain text with encoding detection.
        
        Args:
            file_path: Path to text file to read
        
        Returns:
            ExtractionResult with extracted text and metadata
        """
        import chardet
        
        start_time = time.time()
        
        try:
            # Try multiple encodings
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            
            # Detect encoding
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding') or 'utf-8'
            
            try:
                text = raw_data.decode(encoding)
            except (UnicodeDecodeError, TypeError, LookupError):
                # Fallback to latin-1 which always works
                text = raw_data.decode('latin-1', errors='ignore')
                encoding = 'latin-1'
            
            word_count = len(text.split()) if text else 0
            logger.info(f"Extracted {word_count} words from text file using {encoding}")
            
            return ExtractionResult(
                extracted_text=text,
                word_count=word_count,
                extraction_method=f"text-{encoding}",
                processing_time_ms=(time.time() - start_time) * 1000
            )
            
        except FileNotFoundError:
            logger.error(f"Text file not found: {file_path}")
            return ExtractionResult(
                extracted_text="",
                error=f"Text file not found: {file_path}",
                processing_time_ms=(time.time() - start_time) * 1000
            )
        except Exception as e:
            logger.error(f"Text extraction failed: {e}", exc_info=True)
            return ExtractionResult(
                extracted_text="",
                error=f"Could not read text file: {str(e)}",
                processing_time_ms=(time.time() - start_time) * 1000
            )
